import json
import random
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor

from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.core.exceptions import PermissionDenied
from django.db.models import Q
from django.http import FileResponse, HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.utils import timezone

from .forms import QuestionBankFilterForm, QuestionForm
from .models import (
    Announcement,
    Course,
    Department,
    LandingBackground,
    LandingPanelImage,
    PastQuestion,
    Question,
    QuizAttempt,
    StudentAnswer,
    Textbook,
    get_or_create_student_profile,
)


def home(request):
    if request.user.is_authenticated:
        return redirect('dashboard')
    background_images = LandingBackground.objects.filter(is_active=True)
    hero_panel_image = LandingPanelImage.objects.filter(is_active=True).order_by('order', '-created_at').first()
    return render(request, 'ScholarHub/index.html', {
        'background_images': background_images,
        'hero_panel_image': hero_panel_image,
    })


@login_required
def dashboard(request):
    profile = get_or_create_student_profile(request.user)
    
    # Auto-assign courses based on student profile matching faculty, department, level, semester
    auto_assigned_textbooks = Textbook.objects.select_related('course', 'faculty', 'department').filter(
        is_active=True,
        faculty=profile.faculty,
        department=profile.department,
        level=profile.level,
        semester=profile.semester
    ).distinct().order_by('-created_at')
    
    # Also get elective courses the student manually selected
    elective_textbooks = Textbook.objects.select_related('course', 'faculty', 'department').filter(
        course__in=profile.elective_courses.all(),
        is_active=True
    ).distinct().order_by('-created_at')
    
    # Combine both and get latest past questions
    all_visible_textbooks = (auto_assigned_textbooks | elective_textbooks).distinct().order_by('-created_at')
    recent_textbooks = all_visible_textbooks[:8]
    
    past_questions = PastQuestion.objects.select_related('course', 'faculty', 'department').filter(
        faculty=profile.faculty,
        department=profile.department,
        level=profile.level,
        semester=profile.semester
    ).distinct().order_by('-created_at')[:8]

    context = {
        'profile': profile,
        'auto_assigned_textbooks': auto_assigned_textbooks,
        'recent_textbooks': recent_textbooks,
        'recent_past_questions': past_questions,
        'saved_textbooks': profile.saved_textbooks.all()[:6],
        'saved_past_questions': profile.saved_past_questions.all()[:6],
        'announcements': Announcement.objects.order_by('-created_at')[:5],
    }
    return render(request, 'ScholarHub/dashboard.html', context)


@login_required
def my_courses(request):
    profile = get_or_create_student_profile(request.user)
    
    # Auto-assign textbooks based on student profile
    auto_assigned_textbooks = Textbook.objects.select_related('course', 'faculty', 'department').filter(
        is_active=True,
        faculty=profile.faculty,
        department=profile.department,
        level=profile.level,
        semester=profile.semester
    ).distinct().order_by('course__code')
    
    # Get elective courses the student selected
    elective_textbooks = Textbook.objects.select_related('course', 'faculty', 'department').filter(
        course__in=profile.elective_courses.all(),
        is_active=True
    ).distinct().order_by('course__code')

    context = {
        'profile': profile,
        'auto_assigned_textbooks': auto_assigned_textbooks,
        'elective_textbooks': elective_textbooks,
    }
    return render(request, 'ScholarHub/my_courses.html', context)


@login_required
def textbooks_view(request):
    profile = get_or_create_student_profile(request.user)
    # Show textbooks that are visible to the student based on their profile:
    # - textbooks assigned to their faculty/department/level/semester
    # - textbooks for courses in their profile.courses
    # - textbooks for courses in their elective_courses
    auto_assigned = Textbook.objects.select_related('course', 'faculty', 'department').filter(
        is_active=True,
        faculty=profile.faculty,
        department=profile.department,
        level=profile.level,
        semester=profile.semester
    )
    course_textbooks = Textbook.objects.select_related('course', 'faculty', 'department').filter(
        course__in=profile.courses.all(),
        is_active=True
    )
    elective_textbooks = Textbook.objects.select_related('course', 'faculty', 'department').filter(
        course__in=profile.elective_courses.all(),
        is_active=True
    )

    all_textbooks = (auto_assigned | course_textbooks | elective_textbooks).distinct().order_by('-created_at')

    # Optional free-text search (from global search box fallback)
    q = request.GET.get('q', '').strip()
    if q:
        all_textbooks = all_textbooks.filter(
            Q(title__icontains=q) | Q(author__icontains=q) | Q(course__code__icontains=q)
        )

    context = {
        'profile': profile,
        'all_textbooks': all_textbooks,
    }
    return render(request, 'ScholarHub/textbooks.html', context)


@login_required
def past_questions_view(request):
    profile = get_or_create_student_profile(request.user)
    # Show past questions visible to the student (profile match or course/elective match)
    auto_assigned_pqs = PastQuestion.objects.select_related('course', 'faculty', 'department').filter(
        faculty=profile.faculty,
        department=profile.department,
        level=profile.level,
        semester=profile.semester
    )
    course_pqs = PastQuestion.objects.select_related('course', 'faculty', 'department').filter(
        course__in=profile.courses.all()
    )
    elective_pqs = PastQuestion.objects.select_related('course', 'faculty', 'department').filter(
        course__in=profile.elective_courses.all()
    )

    all_past_questions = (auto_assigned_pqs | course_pqs | elective_pqs).distinct().order_by('-created_at')

    q = request.GET.get('q', '').strip()
    if q:
        all_past_questions = all_past_questions.filter(
            Q(course__code__icontains=q) |
            Q(course__title__icontains=q) |
            Q(description__icontains=q) |
            Q(department__name__icontains=q)
        )

    context = {
        'profile': profile,
        'all_past_questions': all_past_questions,
    }
    return render(request, 'ScholarHub/past_questions.html', context)


@login_required
def saved_books_view(request):
    profile = get_or_create_student_profile(request.user)
    # Filter saved textbooks and past questions to those still visible to the student's profile
    saved_textbooks = profile.saved_textbooks.select_related('course', 'faculty', 'department').filter(
        Q(faculty=profile.faculty, department=profile.department, level=profile.level, semester=profile.semester) |
        Q(course__in=profile.courses.all()) |
        Q(course__in=profile.elective_courses.all())
    ).distinct().order_by('-created_at')

    saved_past_questions = profile.saved_past_questions.select_related('course', 'faculty', 'department').filter(
        Q(faculty=profile.faculty, department=profile.department, level=profile.level, semester=profile.semester) |
        Q(course__in=profile.courses.all()) |
        Q(course__in=profile.elective_courses.all())
    ).distinct().order_by('-created_at')

    context = {
        'profile': profile,
        'saved_textbooks': saved_textbooks,
        'saved_past_questions': saved_past_questions,
    }
    return render(request, 'ScholarHub/saved_books.html', context)


@login_required
def bookmark_textbook(request, pk):
    textbook = get_object_or_404(Textbook, pk=pk)
    profile = get_or_create_student_profile(request.user)
    if profile.saved_textbooks.filter(pk=pk).exists():
        profile.saved_textbooks.remove(textbook)
        messages.info(request, 'Textbook removed from your saved books.')
    else:
        profile.saved_textbooks.add(textbook)
        messages.success(request, 'Textbook saved for later.')
    return redirect(request.META.get('HTTP_REFERER', 'dashboard'))


@login_required
def bookmark_past_question(request, pk):
    question = get_object_or_404(PastQuestion, pk=pk)
    profile = get_or_create_student_profile(request.user)
    if profile.saved_past_questions.filter(pk=pk).exists():
        profile.saved_past_questions.remove(question)
        messages.info(request, 'Past question removed from your saved list.')
    else:
        profile.saved_past_questions.add(question)
        messages.success(request, 'Past question saved for later.')
    return redirect(request.META.get('HTTP_REFERER', 'dashboard'))


@login_required
def textbook_detail(request, pk):
    textbook = get_object_or_404(Textbook, pk=pk)
    profile = get_or_create_student_profile(request.user)
    if not resource_is_visible(profile, textbook):
        raise PermissionDenied('You do not have access to this resource.')
    return render(request, 'ScholarHub/textbook_detail.html', {'resource': textbook, 'profile': profile})


@login_required
def past_question_detail(request, pk):
    question = get_object_or_404(PastQuestion, pk=pk)
    profile = get_or_create_student_profile(request.user)
    if not resource_is_visible(profile, question):
        raise PermissionDenied('You do not have access to this resource.')
    return render(request, 'ScholarHub/past_question_detail.html', {'resource': question, 'profile': profile})


@login_required
def download_textbook(request, pk):
    textbook = get_object_or_404(Textbook, pk=pk)
    profile = get_or_create_student_profile(request.user)
    if not resource_is_visible(profile, textbook):
        raise PermissionDenied('You do not have access to this resource.')
    if not textbook.content:
        raise FileNotFoundError('No note content attached to this textbook.')
    
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading(textbook.title, 0)
    title.runs[0].font.color.rgb = RGBColor(15, 23, 42)  # #0F172A
    
    # Add metadata
    meta = doc.add_paragraph()
    meta.add_run('Course: ').bold = True
    meta.add_run(f'{textbook.course.code} - {textbook.course.title}')
    
    meta = doc.add_paragraph()
    meta.add_run('Author: ').bold = True
    meta.add_run(textbook.author or 'Unknown')
    
    meta = doc.add_paragraph()
    meta.add_run('Department: ').bold = True
    meta.add_run(textbook.department.name)
    
    meta = doc.add_paragraph()
    meta.add_run('Level: ').bold = True
    meta.add_run(textbook.level)
    
    # Add content
    doc.add_heading('Content', level=1)
    doc.add_paragraph(textbook.content)
    
    # Save to BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    response = HttpResponse(doc_io.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{textbook.title.replace(" ", "_")}.docx"'
    return response


@login_required
def view_past_question_pdf(request, pk):
    """View PDF inline in an iframe with SAMEORIGIN framing permission."""
    question = get_object_or_404(PastQuestion, pk=pk)
    profile = get_or_create_student_profile(request.user)
    if not resource_is_visible(profile, question):
        raise PermissionDenied('You do not have access to this resource.')
    if not question.pdf:
        raise FileNotFoundError('No PDF attached to this past question.')
    response = FileResponse(question.pdf.open('rb'), content_type='application/pdf')
    response['X-Frame-Options'] = 'SAMEORIGIN'
    return response


def download_past_question(request, pk):
    question = get_object_or_404(PastQuestion, pk=pk)
    profile = get_or_create_student_profile(request.user)
    if not resource_is_visible(profile, question):
        raise PermissionDenied('You do not have access to this resource.')
    if not question.pdf:
        raise FileNotFoundError('No PDF attached to this past question.')
    response = FileResponse(question.pdf.open('rb'))
    response['Content-Disposition'] = f'attachment; filename="{question.pdf.name.split("/")[-1]}"'
    return response


def departments_for_faculty(request):
    faculty_id = request.GET.get('faculty_id')
    departments = Department.objects.filter(faculty_id=faculty_id).order_by('name') if faculty_id else Department.objects.none()
    payload = [{'id': department.id, 'name': department.name} for department in departments]
    return JsonResponse(payload, safe=False)


def search_api(request):
    q = request.GET.get('q', '').strip()
    results = []
    if q:
        # Textbooks
        t_qs = Textbook.objects.select_related('course').filter(
            Q(title__icontains=q) | Q(author__icontains=q) | Q(course__code__icontains=q)
        )[:6]
        for t in t_qs:
            results.append({
                'type': 'Textbook',
                'label': f'{t.title} — {t.course.code}',
                'url': reverse('textbook_detail', args=[t.pk])
            })

        # Courses
        c_qs = Course.objects.filter(Q(code__icontains=q) | Q(title__icontains=q))[:6]
        for c in c_qs:
            results.append({
                'type': 'Course',
                'label': f'{c.code} — {c.title}',
                'url': reverse('my_courses') + f'#course-{c.id}'
            })

        # Past questions
        p_qs = PastQuestion.objects.select_related('course').filter(Q(course__code__icontains=q) | Q(academic_session__icontains=q))[:6]
        for p in p_qs:
            results.append({
                'type': 'PastQuestion',
                'label': f'{p.course.code} — {p.academic_session}',
                'url': reverse('past_question_detail', args=[p.pk])
            })

    return JsonResponse({'results': results})


def question_bank_admin_view(request):
    if not request.user.is_staff:
        return redirect('dashboard')

    if request.method == 'POST' and request.POST.get('delete_question'):
        question_id = request.POST.get('delete_question')
        Question.objects.filter(pk=question_id).delete()
        messages.success(request, 'Question deleted successfully.')
        return redirect('question_bank_admin')

    edit_id = request.GET.get('edit_question')
    question_instance = None
    if edit_id:
        question_instance = get_object_or_404(Question, pk=edit_id)

    filter_form = QuestionBankFilterForm(request.GET or None)
    questions = Question.objects.select_related('course', 'faculty', 'department').all()

    if filter_form.is_valid():
        faculty = filter_form.cleaned_data.get('faculty')
        department = filter_form.cleaned_data.get('department')
        level = filter_form.cleaned_data.get('level')
        semester = filter_form.cleaned_data.get('semester')
        course = filter_form.cleaned_data.get('course')

        if faculty:
            questions = questions.filter(faculty=faculty)
        if department:
            questions = questions.filter(department=department)
        if level:
            questions = questions.filter(level=level)
        if semester:
            questions = questions.filter(semester=semester)
        if course:
            questions = questions.filter(course=course)

    search_query = request.GET.get('q', '').strip()
    if search_query:
        questions = questions.filter(Q(text__icontains=search_query) | Q(course__code__icontains=search_query) | Q(course__title__icontains=search_query))

    if request.method == 'POST':
        question_form = QuestionForm(request.POST, instance=question_instance)
        if question_form.is_valid():
            question_form.save()
            messages.success(request, 'Question saved successfully.')
            return redirect('question_bank_admin')
    else:
        question_form = QuestionForm(instance=question_instance)

    context = {
        'filter_form': filter_form,
        'question_form': question_form,
        'questions': questions.order_by('course__code', 'text'),
        'selected_course': filter_form.cleaned_data.get('course') if filter_form.is_valid() else None,
        'editing_question': question_instance,
    }
    return render(request, 'ScholarHub/question_bank_admin.html', context)


@login_required
def start_quiz(request, course_id):
    profile = get_or_create_student_profile(request.user)
    course = get_object_or_404(Course, pk=course_id)
    if not resource_is_visible(profile, course):
        raise PermissionDenied('You do not have access to this course.')

    questions = list(Question.objects.filter(course=course).order_by('?'))
    if not questions:
        messages.info(request, 'No questions are available for this course yet.')
        return redirect('past_questions')

    request.session['quiz_course_id'] = course.id
    request.session['quiz_question_ids'] = [question.id for question in questions]
    request.session['quiz_answers'] = {}
    request.session['quiz_started_at'] = timezone.now().timestamp()
    request.session['quiz_time_limit'] = course.quiz_time_limit or 0
    return render(request, 'ScholarHub/quiz_start.html', {'course': course, 'questions': questions, 'profile': profile})


@login_required
def quiz_question_view(request, course_id):
    profile = get_or_create_student_profile(request.user)
    course = get_object_or_404(Course, pk=course_id)
    if not resource_is_visible(profile, course):
        raise PermissionDenied('You do not have access to this course.')

    question_ids = request.session.get('quiz_question_ids', [])
    if not question_ids:
        return redirect('past_questions')

    index = int(request.GET.get('index', 0))
    if index < 0:
        index = 0
    if index >= len(question_ids):
        index = len(question_ids) - 1

    question = Question.objects.get(pk=question_ids[index])
    options = list(question.get_options().items())
    random.shuffle(options)
    current_time_limit = request.session.get('quiz_time_limit', 0) or 0
    time_remaining = None
    if current_time_limit:
        started = request.session.get('quiz_started_at')
        if started:
            elapsed_seconds = int(timezone.now().timestamp() - started)
            time_remaining = max(0, current_time_limit * 60 - elapsed_seconds)
    saved_answer = request.session.get('quiz_answers', {}).get(str(question.id), '').upper()
    return render(request, 'ScholarHub/quiz_page.html', {
        'course': course,
        'question': question,
        'options': options,
        'index': index,
        'question_count': len(question_ids),
        'profile': profile,
        'time_remaining': time_remaining,
        'time_limit': current_time_limit,
        'saved_answer': saved_answer,
    })


@login_required
def save_quiz_answer(request, course_id):
    profile = get_or_create_student_profile(request.user)
    course = get_object_or_404(Course, pk=course_id)
    if not resource_is_visible(profile, course):
        raise PermissionDenied('You do not have access to this course.')

    if request.method == 'POST':
        question_id = request.POST.get('question_id')
        answer = (request.POST.get('answer') or '').upper()
        if question_id:
            answers = request.session.get('quiz_answers', {})
            answers[question_id] = answer
            request.session['quiz_answers'] = answers
    return JsonResponse({'status': 'ok'})


@login_required
def submit_quiz(request, course_id):
    profile = get_or_create_student_profile(request.user)
    course = get_object_or_404(Course, pk=course_id)
    if not resource_is_visible(profile, course):
        raise PermissionDenied('You do not have access to this course.')

    question_ids = request.session.get('quiz_question_ids', [])
    answers = request.session.get('quiz_answers', {})
    for key, value in request.POST.items():
        if key.startswith('question_'):
            question_id = key.split('_', 1)[1]
            answers[question_id] = value
    request.session['quiz_answers'] = answers

    if not question_ids:
        return redirect('past_questions')

    questions = list(Question.objects.filter(pk__in=question_ids))
    score = 0
    student_answers = []
    for question in questions:
        selected_answer = answers.get(str(question.id), '').upper()
        is_correct = selected_answer == question.correct_answer
        if is_correct:
            score += 1
        student_answers.append((question, selected_answer, is_correct))

    total_questions = len(questions)
    percentage = round((score / total_questions) * 100, 2) if total_questions else 0.0
    time_taken = 0
    started_at = request.session.get('quiz_started_at')
    auto_submitted = False
    if started_at:
        time_taken = int(timezone.now().timestamp() - started_at)
        if course.quiz_time_limit and time_taken >= course.quiz_time_limit * 60:
            auto_submitted = True

    attempt = QuizAttempt.objects.create(
        student=profile,
        course=course,
        score=score,
        percentage=percentage,
        total_questions=total_questions,
        correct_answers=score,
        wrong_answers=total_questions - score,
        time_taken=time_taken,
        time_limit=course.quiz_time_limit,
        auto_submitted=auto_submitted,
    )

    for question, selected_answer, is_correct in student_answers:
        StudentAnswer.objects.create(
            quiz_attempt=attempt,
            question=question,
            selected_answer=selected_answer,
            correct_answer=question.correct_answer,
            is_correct=is_correct,
        )

    request.session.pop('quiz_question_ids', None)
    request.session.pop('quiz_answers', None)
    request.session.pop('quiz_started_at', None)
    request.session.pop('quiz_time_limit', None)
    request.session['quiz_review_attempt_id'] = attempt.id
    return render(request, 'ScholarHub/quiz_results.html', {'attempt': attempt, 'questions': questions, 'student_answers': student_answers, 'course': course, 'auto_submitted': auto_submitted})


@login_required
def quiz_review(request, attempt_id):
    attempt = get_object_or_404(QuizAttempt, pk=attempt_id, student__user=request.user)
    student_answers = list(attempt.student_answers.select_related('question').all())
    return render(request, 'ScholarHub/quiz_review.html', {'attempt': attempt, 'student_answers': student_answers})


@login_required
def quiz_attempt_history(request):
    profile = get_or_create_student_profile(request.user)
    attempts = QuizAttempt.objects.filter(student=profile).select_related('course').order_by('-date_completed')
    return render(request, 'ScholarHub/quiz_attempt_history.html', {'attempts': attempts, 'profile': profile})


def resource_is_visible(profile, resource):
    if profile.faculty and profile.department and getattr(resource, 'faculty_id', None) == profile.faculty_id and getattr(resource, 'department_id', None) == profile.department_id and getattr(resource, 'level', None) == profile.level:
        return True

    if isinstance(resource, Course):
        if getattr(resource, 'department', None) and getattr(resource.department, 'faculty_id', None) == profile.faculty_id and getattr(resource.department, 'id', None) == profile.department_id and getattr(resource, 'level', None) == profile.level:
            return True

    course_ids = set(profile.courses.values_list('id', flat=True)) | set(profile.elective_courses.values_list('id', flat=True))
    return getattr(resource, 'course_id', None) in course_ids


def custom_permission_denied_view(request, exception=None):
    from django.shortcuts import render
    return render(request, 'ScholarHub/error.html', {'error_message': str(exception) if exception else 'Permission denied.'}, status=403)

def custom_page_not_found_view(request, exception=None):
    from django.shortcuts import render
    return render(request, 'ScholarHub/error.html', {'error_message': 'Page not found.'}, status=404)

def custom_server_error_view(request):
    from django.shortcuts import render
    return render(request, 'ScholarHub/error.html', {'error_message': 'An unexpected error occurred. Please try again later.'}, status=500)
